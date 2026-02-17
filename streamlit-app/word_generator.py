# -*- coding: utf-8 -*-
"""
Word出力モジュール (v3.0)
プレスリリースのWord文書を生成

テンプレート形式: {{KEY}} プレースホルダー方式
使用テンプレート: _archive/テンプレート/【テンプレ】プレスリリース_v4.docx

v3.0 新機能:
- 評価項目別ランキング表の動的挿入
- 部門別ランキング表の動的挿入
- 企業別プレースホルダー対応 ({{RANK_1_COMPANY}}, {{RANK_1_SCORE}})
- 表スタイリング強化 (ヘッダー色、セル幅、フォント設定)
- テンプレート検証機能
- 複数テンプレート対応
"""

import os
import re
import logging
from io import BytesIO
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

logger = logging.getLogger(__name__)

# テンプレートパス
TEMPLATE_DIR = os.path.join(
    os.path.dirname(os.path.dirname(__file__)),
    "_archive", "テンプレート"
)

# 利用可能なテンプレート
TEMPLATES = {
    "v3": os.path.join(TEMPLATE_DIR, "【テンプレ】プレスリリース_v3.docx"),
    "v4": os.path.join(TEMPLATE_DIR, "【テンプレ】プレスリリース_v4.docx"),
}

# デフォルトテンプレート
DEFAULT_TEMPLATE = "v4"

# 曜日マッピング
WEEKDAY_JP = ["月", "火", "水", "木", "金", "土", "日"]

# 表スタイル設定
TABLE_STYLES = {
    "header_bg_color": "1F4E79",  # 濃い青
    "header_text_color": "FFFFFF",  # 白
    "alt_row_color": "F2F2F2",  # 薄いグレー（交互行）
    "border_color": "B4B4B4",  # ボーダー色
    "font_name": "游ゴシック",
    "font_size": 10,
    "header_font_size": 10,
}


class WordGenerator:
    """Wordプレスリリース生成クラス (v3.0 - 複数表・スタイリング対応)"""

    def __init__(
        self,
        ranking_name: str,
        year: int,
        month: int = None,
        day: int = None,
        template_version: str = None
    ):
        """
        Args:
            ranking_name: ランキング名（例: "ネット証券"）
            year: 発表年度
            month: 発表月（デフォルト: 現在月）
            day: 発表日（デフォルト: 現在日）
            template_version: テンプレートバージョン ("v3" or "v4")
        """
        self.ranking_name = ranking_name
        self.year = year
        self.month = month or datetime.now().month
        self.day = day or datetime.now().day
        self.doc = None
        self.template_version = template_version or DEFAULT_TEMPLATE

        # 日付計算
        try:
            dt = datetime(year, self.month, self.day)
            self.weekday = WEEKDAY_JP[dt.weekday()]
        except (ValueError, TypeError):
            self.weekday = ""

    def get_template_path(self) -> str:
        """テンプレートパスを取得"""
        path = TEMPLATES.get(self.template_version)
        if path and os.path.exists(path):
            return path
        # フォールバック: v3を試す
        fallback = TEMPLATES.get("v3")
        if fallback and os.path.exists(fallback):
            logger.warning(f"テンプレート {self.template_version} が見つかりません。v3 にフォールバック")
            return fallback
        return None

    def load_template(self, template_path: str = None) -> bool:
        """テンプレートを読み込む"""
        path = template_path or self.get_template_path()

        if not path:
            logger.error("利用可能なテンプレートがありません")
            return False

        if not os.path.exists(path):
            logger.error(f"テンプレートが見つかりません: {path}")
            return False

        try:
            self.doc = Document(path)
            logger.info(f"テンプレート読み込み成功: {path}")
            return True
        except Exception as e:
            logger.error(f"テンプレート読み込みエラー: {e}")
            return False

    def validate_template(self) -> Tuple[bool, List[str], List[str]]:
        """テンプレートのプレースホルダーを検証

        Returns:
            (is_valid, found_placeholders, missing_placeholders)
        """
        if not self.doc:
            return False, [], ["テンプレートが読み込まれていません"]

        required = {
            "DATE", "YEAR", "WEEKDAY", "RANKING_NAME", "HEADLINE",
            "SUBHEADLINE", "TOPIC_1", "TOPIC_1_DETAIL"
        }

        found = set()

        # 段落を検索
        for para in self.doc.paragraphs:
            matches = re.findall(r'\{\{([^}]+)\}\}', para.text)
            found.update(matches)

        # テーブルを検索
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        matches = re.findall(r'\{\{([^}]+)\}\}', para.text)
                        found.update(matches)

        missing = required - found
        is_valid = len(missing) == 0

        return is_valid, list(found), list(missing)

    def _replace_in_paragraph(self, para, replacements: Dict[str, str]):
        """段落内のプレースホルダーを置換"""
        text = para.text
        modified = False

        for key, value in replacements.items():
            placeholder = f"{{{{{key}}}}}"  # {{KEY}} 形式
            if placeholder in text:
                text = text.replace(placeholder, str(value) if value else "")
                modified = True

        if modified and para.runs:
            # 最初のRunに全テキストを設定（書式保持）
            for run in para.runs:
                run.text = ""
            para.runs[0].text = text

    def _set_cell_shading(self, cell, color: str):
        """セルの背景色を設定"""
        shading_elm = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{color}"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def _set_cell_text_color(self, cell, color: str):
        """セルのテキスト色を設定"""
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor.from_string(color)

    def _create_styled_table(
        self,
        data: List[List[str]],
        headers: List[str],
        col_widths: List[float] = None,
        title: str = None
    ):
        """スタイル付きテーブルを作成

        Args:
            data: データ行のリスト
            headers: ヘッダー行
            col_widths: 各列の幅（cm）
            title: 表のタイトル

        Returns:
            作成したテーブル
        """
        if not self.doc:
            return None

        # タイトルを追加
        if title:
            title_para = self.doc.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.bold = True
            title_run.font.size = Pt(11)
            title_para.space_after = Pt(6)

        # テーブル作成
        num_cols = len(headers)
        table = self.doc.add_table(rows=1, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # ヘッダー行
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            cell = header_cells[i]
            cell.text = header
            # ヘッダースタイル
            self._set_cell_shading(cell, TABLE_STYLES["header_bg_color"])
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(TABLE_STYLES["header_font_size"])
                    run.font.color.rgb = RGBColor.from_string(TABLE_STYLES["header_text_color"])

        # データ行
        for row_idx, row_data in enumerate(data):
            row_cells = table.add_row().cells
            for col_idx, value in enumerate(row_data):
                cell = row_cells[col_idx]
                cell.text = str(value) if value is not None else ""
                # 交互行の背景色
                if row_idx % 2 == 1:
                    self._set_cell_shading(cell, TABLE_STYLES["alt_row_color"])
                # セルのテキストスタイル
                for para in cell.paragraphs:
                    if col_idx == 0:  # 順位列は中央揃え
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif col_idx == len(row_data) - 1:  # 得点列は右揃え
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for run in para.runs:
                        run.font.size = Pt(TABLE_STYLES["font_size"])

        # 列幅設定
        if col_widths:
            for row in table.rows:
                for i, width in enumerate(col_widths):
                    if i < len(row.cells):
                        row.cells[i].width = Cm(width)

        # 空行を追加
        self.doc.add_paragraph()

        return table

    def replace_placeholders(
        self,
        overall_data: List[Dict] = None,
        topics: List[str] = None,
        topic_details: List[str] = None,
        highlights: List[str] = None,
        subheadline: str = None,
        sample_size: int = None,
        min_sample: int = 100,
        company_count: int = None,
        ranking_url: str = None,
        prev_year_data: List[Dict] = None
    ):
        """プレースホルダーを置換

        Args:
            overall_data: 総合ランキングデータ
            topics: TOPICSリスト（タイトル）
            topic_details: TOPICS詳細リスト
            highlights: ハイライト（見出し）リスト
            subheadline: サブ見出し
            sample_size: 回答者数
            min_sample: 規定人数
            company_count: 調査企業数
            ranking_url: ランキングページURL
            prev_year_data: 前年データ（変動表示用）
        """
        if not self.doc:
            logger.error("テンプレートが読み込まれていません")
            return

        # 基本置換マッピング
        replacements = {
            # 日付関連
            "DATE": f"{self.year}年{self.month}月{self.day}日",
            "YEAR": f"{self.year}年",
            "WEEKDAY": self.weekday,
            "RELEASE_DATE_SLASH": f"{self.year}/{self.month:02d}/{self.day:02d}",

            # ランキング情報
            "RANKING_NAME": self.ranking_name,

            # 見出し
            "HEADLINE": highlights[0] if highlights else "",
            "SUBHEADLINE": subheadline or "",

            # TOPICS (3セット)
            "TOPIC_1": topics[0] if topics and len(topics) > 0 else "",
            "TOPIC_1_DETAIL": topic_details[0] if topic_details and len(topic_details) > 0 else "",
            "TOPIC_2": topics[1] if topics and len(topics) > 1 else "",
            "TOPIC_2_DETAIL": topic_details[1] if topic_details and len(topic_details) > 1 else "",
            "TOPIC_3": topics[2] if topics and len(topics) > 2 else "",
            "TOPIC_3_DETAIL": topic_details[2] if topic_details and len(topic_details) > 2 else "",

            # 調査概要
            "SAMPLE_SIZE": f"{sample_size:,}" if sample_size else "",
            "MIN_SAMPLE": str(min_sample),
            "COMPANY_COUNT": str(company_count) if company_count else "",
            "RANKING_URL": ranking_url or "",
        }

        # 企業別プレースホルダー (TOP10)
        if overall_data:
            for i, entry in enumerate(overall_data[:10], 1):
                replacements[f"RANK_{i}_COMPANY"] = entry.get("company", "")
                score = entry.get("score")
                replacements[f"RANK_{i}_SCORE"] = f"{score:.1f}" if score is not None else ""
                replacements[f"RANK_{i}_RANK"] = str(entry.get("rank", i))

                # 前年比（前年データがある場合）
                if prev_year_data:
                    company = entry.get("company", "")
                    prev_entry = next((e for e in prev_year_data if e.get("company") == company), None)
                    if prev_entry:
                        prev_rank = prev_entry.get("rank")
                        curr_rank = entry.get("rank", i)
                        if prev_rank and curr_rank:
                            diff = prev_rank - curr_rank
                            if diff > 0:
                                replacements[f"RANK_{i}_CHANGE"] = f"↑{diff}"
                            elif diff < 0:
                                replacements[f"RANK_{i}_CHANGE"] = f"↓{abs(diff)}"
                            else:
                                replacements[f"RANK_{i}_CHANGE"] = "→"
                    else:
                        replacements[f"RANK_{i}_CHANGE"] = "NEW"
                else:
                    replacements[f"RANK_{i}_CHANGE"] = ""

        # 全段落に対して置換を実行
        for para in self.doc.paragraphs:
            self._replace_in_paragraph(para, replacements)

        # テーブル内のセルも置換
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph(para, replacements)

    def add_overall_ranking_table(
        self,
        overall_data: List[Dict],
        top_n: int = 10,
        title: str = None,
        show_score: bool = True
    ):
        """総合ランキング表を追加

        Args:
            overall_data: ランキングデータ
            top_n: 表示件数
            title: 表のタイトル
            show_score: 得点を表示するか
        """
        if not self.doc or not overall_data:
            return

        headers = ["順位", "企業名"]
        col_widths = [1.5, 6.0]

        if show_score:
            headers.append("得点")
            col_widths.append(2.0)

        data = []
        for entry in overall_data[:top_n]:
            row = [
                f"{entry.get('rank', '-')}位",
                entry.get('company', '')
            ]
            if show_score:
                score = entry.get('score')
                row.append(f"{score:.2f}点" if score is not None else "-")
            data.append(row)

        self._create_styled_table(
            data=data,
            headers=headers,
            col_widths=col_widths,
            title=title or f"【{self.ranking_name}】総合ランキング TOP{top_n}"
        )

    def add_item_ranking_tables(
        self,
        item_data: Dict[str, List[Dict]],
        top_n: int = 5,
        selected_items: List[str] = None
    ):
        """評価項目別ランキング表を追加

        Args:
            item_data: 評価項目別データ {項目名: [ランキングデータ]}
            top_n: 各項目の表示件数
            selected_items: 表示する項目（Noneで全項目）
        """
        if not self.doc or not item_data:
            return

        items_to_show = selected_items or list(item_data.keys())

        for item_name in items_to_show:
            if item_name not in item_data:
                continue

            rankings = item_data[item_name]
            if not rankings:
                continue

            headers = ["順位", "企業名", "得点"]
            col_widths = [1.5, 5.5, 2.0]

            data = []
            for entry in rankings[:top_n]:
                row = [
                    f"{entry.get('rank', '-')}位",
                    entry.get('company', ''),
                    f"{entry.get('score', 0):.2f}点"
                ]
                data.append(row)

            self._create_styled_table(
                data=data,
                headers=headers,
                col_widths=col_widths,
                title=f"【評価項目】{item_name} TOP{top_n}"
            )

    def add_dept_ranking_tables(
        self,
        dept_data: Dict[str, Dict],
        top_n: int = 5,
        selected_depts: List[str] = None
    ):
        """部門別ランキング表を追加

        Args:
            dept_data: 部門別データ {部門名: {year: [ランキングデータ]}}
            top_n: 各部門の表示件数
            selected_depts: 表示する部門（Noneで全部門）
        """
        if not self.doc or not dept_data:
            return

        depts_to_show = selected_depts or list(dept_data.keys())

        for dept_name in depts_to_show:
            if dept_name not in dept_data:
                continue

            dept_years = dept_data[dept_name]
            if not dept_years:
                continue

            # 最新年のデータを使用
            latest_year = max(dept_years.keys()) if dept_years else None
            if not latest_year:
                continue

            rankings = dept_years.get(latest_year, [])
            if not rankings:
                continue

            headers = ["順位", "企業名", "得点"]
            col_widths = [1.5, 5.5, 2.0]

            data = []
            for entry in rankings[:top_n]:
                row = [
                    f"{entry.get('rank', '-')}位",
                    entry.get('company', ''),
                    f"{entry.get('score', 0):.2f}点"
                ]
                data.append(row)

            self._create_styled_table(
                data=data,
                headers=headers,
                col_widths=col_widths,
                title=f"【部門別】{dept_name} TOP{top_n}"
            )

    def add_comparison_table(
        self,
        current_data: List[Dict],
        prev_data: List[Dict],
        top_n: int = 10,
        title: str = None
    ):
        """前年比較表を追加

        Args:
            current_data: 今年のランキングデータ
            prev_data: 前年のランキングデータ
            top_n: 表示件数
            title: 表のタイトル
        """
        if not self.doc or not current_data:
            return

        headers = ["順位", "企業名", "得点", "前年順位", "変動"]
        col_widths = [1.2, 4.5, 1.5, 1.5, 1.3]

        # 前年データをルックアップ用に変換
        prev_lookup = {}
        if prev_data:
            for entry in prev_data:
                company = entry.get("company", "")
                if company:
                    prev_lookup[company] = entry

        data = []
        for entry in current_data[:top_n]:
            company = entry.get("company", "")
            curr_rank = entry.get("rank", "-")
            score = entry.get("score")

            prev_entry = prev_lookup.get(company)
            if prev_entry:
                prev_rank = prev_entry.get("rank", "-")
                if isinstance(curr_rank, int) and isinstance(prev_rank, int):
                    diff = prev_rank - curr_rank
                    if diff > 0:
                        change = f"↑{diff}"
                    elif diff < 0:
                        change = f"↓{abs(diff)}"
                    else:
                        change = "→"
                else:
                    change = "-"
            else:
                prev_rank = "-"
                change = "NEW"

            row = [
                f"{curr_rank}位",
                company,
                f"{score:.2f}点" if score is not None else "-",
                f"{prev_rank}位" if prev_rank != "-" else "-",
                change
            ]
            data.append(row)

        self._create_styled_table(
            data=data,
            headers=headers,
            col_widths=col_widths,
            title=title or f"【{self.ranking_name}】前年比較 TOP{top_n}"
        )

    def generate(
        self,
        overall_data: List[Dict] = None,
        topics: List[str] = None,
        topic_details: List[str] = None,
        highlights: List[str] = None,
        subheadline: str = None,
        sample_size: int = None,
        min_sample: int = 100,
        company_count: int = None,
        ranking_url: str = None,
        include_overall_table: bool = False,
        include_item_tables: bool = False,
        include_dept_tables: bool = False,
        include_comparison_table: bool = False,
        item_data: Dict = None,
        dept_data: Dict = None,
        prev_year_data: List[Dict] = None,
        selected_items: List[str] = None,
        selected_depts: List[str] = None,
        table_top_n: int = 10
    ) -> Optional[BytesIO]:
        """Word文書を生成

        Args:
            overall_data: 総合ランキングデータ
            topics: TOPICSリスト
            topic_details: TOPICS詳細リスト
            highlights: ハイライトリスト
            subheadline: サブ見出し
            sample_size: 回答者数
            min_sample: 規定人数
            company_count: 調査企業数
            ranking_url: ランキングページURL
            include_overall_table: 総合ランキング表を含めるか
            include_item_tables: 評価項目別表を含めるか
            include_dept_tables: 部門別表を含めるか
            include_comparison_table: 前年比較表を含めるか
            item_data: 評価項目別データ
            dept_data: 部門別データ
            prev_year_data: 前年データ
            selected_items: 表示する評価項目
            selected_depts: 表示する部門
            table_top_n: 各表の表示件数

        Returns:
            BytesIOオブジェクト（Wordファイル）
        """
        # テンプレート読み込み
        if not self.load_template():
            return None

        # プレースホルダー置換
        self.replace_placeholders(
            overall_data=overall_data,
            topics=topics,
            topic_details=topic_details,
            highlights=highlights,
            subheadline=subheadline,
            sample_size=sample_size,
            min_sample=min_sample,
            company_count=company_count,
            ranking_url=ranking_url,
            prev_year_data=prev_year_data
        )

        # 表を追加（オプション）
        if include_overall_table and overall_data:
            self.add_overall_ranking_table(overall_data, top_n=table_top_n)

        if include_comparison_table and overall_data and prev_year_data:
            self.add_comparison_table(
                current_data=overall_data,
                prev_data=prev_year_data,
                top_n=table_top_n
            )

        if include_item_tables and item_data:
            self.add_item_ranking_tables(
                item_data=item_data,
                top_n=min(5, table_top_n),
                selected_items=selected_items
            )

        if include_dept_tables and dept_data:
            self.add_dept_ranking_tables(
                dept_data=dept_data,
                top_n=min(5, table_top_n),
                selected_depts=selected_depts
            )

        # BytesIOに保存
        output = BytesIO()
        self.doc.save(output)
        output.seek(0)

        return output


def generate_word_release(
    ranking_name: str,
    year: int,
    overall_data: List[Dict] = None,
    topics: List[str] = None,
    topic_details: List[str] = None,
    highlights: List[str] = None,
    subheadline: str = None,
    month: int = None,
    day: int = None,
    sample_size: int = None,
    company_count: int = None,
    ranking_url: str = None,
    include_table: bool = False,
    include_overall_table: bool = None,
    include_item_tables: bool = False,
    include_dept_tables: bool = False,
    include_comparison_table: bool = False,
    item_data: Dict = None,
    dept_data: Dict = None,
    prev_year_data: List[Dict] = None,
    selected_items: List[str] = None,
    selected_depts: List[str] = None,
    table_top_n: int = 10,
    template_version: str = None
) -> Optional[BytesIO]:
    """Wordプレスリリースを生成（簡易インターフェース）

    Args:
        ranking_name: ランキング名
        year: 発表年度
        overall_data: 総合ランキングデータ
        topics: TOPICSリスト
        topic_details: TOPICS詳細リスト
        highlights: ハイライトリスト
        subheadline: サブ見出し
        month: 発表月
        day: 発表日
        sample_size: 回答者数
        company_count: 調査企業数
        ranking_url: ランキングページURL
        include_table: 表を含めるか（後方互換性のため）
        include_overall_table: 総合ランキング表を含めるか
        include_item_tables: 評価項目別表を含めるか
        include_dept_tables: 部門別表を含めるか
        include_comparison_table: 前年比較表を含めるか
        item_data: 評価項目別データ
        dept_data: 部門別データ
        prev_year_data: 前年データ
        selected_items: 表示する評価項目
        selected_depts: 表示する部門
        table_top_n: 表の表示件数
        template_version: テンプレートバージョン

    Returns:
        BytesIOオブジェクト
    """
    generator = WordGenerator(
        ranking_name=ranking_name,
        year=year,
        month=month,
        day=day,
        template_version=template_version
    )

    # 後方互換性: include_table が指定されていて include_overall_table が未指定の場合
    if include_overall_table is None:
        include_overall_table = include_table

    return generator.generate(
        overall_data=overall_data,
        topics=topics,
        topic_details=topic_details,
        highlights=highlights,
        subheadline=subheadline,
        sample_size=sample_size,
        company_count=company_count,
        ranking_url=ranking_url,
        include_overall_table=include_overall_table,
        include_item_tables=include_item_tables,
        include_dept_tables=include_dept_tables,
        include_comparison_table=include_comparison_table,
        item_data=item_data,
        dept_data=dept_data,
        prev_year_data=prev_year_data,
        selected_items=selected_items,
        selected_depts=selected_depts,
        table_top_n=table_top_n
    )


def get_available_templates() -> Dict[str, str]:
    """利用可能なテンプレート一覧を取得

    Returns:
        {バージョン: ファイルパス} の辞書（存在するテンプレートのみ）
    """
    available = {}
    for version, path in TEMPLATES.items():
        if os.path.exists(path):
            available[version] = path
    return available


def validate_template_file(template_path: str) -> Tuple[bool, List[str], List[str]]:
    """テンプレートファイルを検証

    Args:
        template_path: テンプレートのパス

    Returns:
        (is_valid, found_placeholders, warnings)
    """
    try:
        doc = Document(template_path)
    except Exception as e:
        return False, [], [f"ファイルを開けません: {e}"]

    required = {"DATE", "RANKING_NAME", "HEADLINE"}
    found = set()
    warnings = []

    for para in doc.paragraphs:
        matches = re.findall(r'\{\{([^}]+)\}\}', para.text)
        found.update(matches)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    matches = re.findall(r'\{\{([^}]+)\}\}', para.text)
                    found.update(matches)

    missing = required - found
    if missing:
        warnings.append(f"必須プレースホルダーが不足: {', '.join(missing)}")

    is_valid = len(missing) == 0
    return is_valid, list(found), warnings


# ========================================
# デバッグ用
# ========================================
if __name__ == "__main__":
    import sys
    sys.stdout.reconfigure(encoding='utf-8')

    print("=== Word Generator v3.0 Test ===")

    # 利用可能なテンプレートを表示
    available = get_available_templates()
    print(f"利用可能なテンプレート: {list(available.keys())}")

    # テストデータ
    test_data = [
        {"rank": 1, "company": "SBI証券", "score": 68.9},
        {"rank": 1, "company": "楽天証券", "score": 68.9},
        {"rank": 3, "company": "マネックス証券", "score": 67.5},
        {"rank": 4, "company": "松井証券", "score": 66.8},
        {"rank": 5, "company": "auカブコム証券", "score": 65.2},
    ]

    prev_data = [
        {"rank": 1, "company": "SBI証券", "score": 68.5},
        {"rank": 2, "company": "楽天証券", "score": 68.0},
        {"rank": 4, "company": "マネックス証券", "score": 66.0},
        {"rank": 3, "company": "松井証券", "score": 67.0},
    ]

    test_item_data = {
        "取引手数料": [
            {"rank": 1, "company": "SBI証券", "score": 72.5},
            {"rank": 2, "company": "楽天証券", "score": 71.0},
            {"rank": 3, "company": "松井証券", "score": 69.5},
        ],
        "取引ツールの使いやすさ": [
            {"rank": 1, "company": "楽天証券", "score": 70.8},
            {"rank": 2, "company": "SBI証券", "score": 69.2},
            {"rank": 3, "company": "マネックス証券", "score": 68.5},
        ],
    }

    test_topics = [
        "SBI証券と楽天証券が同率1位",
        "SBI証券が3年連続1位を達成",
        "マネックス証券が初のTOP3入り"
    ]

    test_topic_details = [
        "今年度の調査では、SBI証券と楽天証券が68.9点で並び、初の同率1位となりました。",
        "SBI証券は2024年から3年連続で総合1位を獲得。取引手数料の評価が特に高い結果となりました。",
        "マネックス証券は前年5位から3位に躍進。投資情報の充実度が評価されました。"
    ]

    test_highlights = ["SBI証券と楽天証券が同率1位、3年連続の快挙"]

    # 生成テスト
    result = generate_word_release(
        ranking_name="ネット証券",
        year=2026,
        month=1,
        day=15,
        overall_data=test_data,
        topics=test_topics,
        topic_details=test_topic_details,
        highlights=test_highlights,
        subheadline="業界初の同率1位、手数料競争が加速",
        sample_size=5000,
        company_count=15,
        ranking_url="https://cs.oricon.co.jp/rank/netsec/",
        include_overall_table=True,
        include_comparison_table=True,
        include_item_tables=True,
        prev_year_data=prev_data,
        item_data=test_item_data,
        table_top_n=5,
        template_version="v3"  # v4がなければv3を使用
    )

    if result:
        output_path = "test_release_v3.docx"
        with open(output_path, "wb") as f:
            f.write(result.getvalue())
        print(f"[OK] Generated: {output_path}")
    else:
        print("[ERROR] Generation failed")
